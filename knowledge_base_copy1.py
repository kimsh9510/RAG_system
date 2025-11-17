#벡터 DB구축 - > 지식그래프 DB로 변경 예정
from PyPDF2 import PdfReader
import subprocess
import sys
import pdfplumber
from langchain.schema import Document
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from models import load_embeddings
import glob
import docx # python-docx
import os
import zipfile #hwpx 
import xml.etree.ElementTree as ET #hwxp
import pandas as pd
import re
import json
import subprocess
import sys

def _load_geojson_as_docs(file_path: str):
    """
    Load a .geojson file and convert each Feature's properties into a LangChain Document.
    - page_content: a single JSON object that contains
        * the original properties from GeoJSON
        * additional GIS_인구 risk fields
    """
    # If the caller requested a geojson file, run the generator unconditionally so
    # the latest location query output is produced.
    if file_path.lower().endswith('.geojson'):
        try:
            script_path = os.path.join(os.path.dirname(__file__), "Process_GIS_Data.py")
            if os.path.exists(script_path):
                print(f"Running generator for requested geojson: {script_path}")
                proc = subprocess.run(
                    [sys.executable, script_path],
                    capture_output=True,
                    text=True,
                    check=False
                )
                if proc.returncode != 0:
                    print(f"Process_GIS_Data.py exited with code {proc.returncode}. stderr:\n{proc.stderr}")
                else:
                    print("Process_GIS_Data.py ran successfully. stdout:", proc.stdout)
            else:
                print(f"Generator script not found at: {script_path}")
        except Exception as e:
            print(f"Failed to run Process_GIS_Data.py: {e}")

    # After attempting generation, ensure the geojson exists
    if not os.path.exists(file_path):
        print(f"[ERROR] GeoJSON file still missing after generation attempt: {file_path}")
        return []

    # ----- helpers -----
    # Safely turn any value into a float without crashing.
    def _safe_float(x, default=0.0):
        try:
            if x is None:
                return float(default)
            if isinstance(x, str) and x.strip() == "":
                return float(default)
            return float(x)
        except Exception:
            return float(default)

    # Safely convert a value to a clean string.
    def _norm_str(x):
        return (x or "").strip()
    
    # A set of 32 characters used for geohash encoding.
    _BASE32 = "0123456789bcdefghjkmnpqrstuvwxyz"

    # Simple geohash implementation turning (lat, lon) 
    # into a short string like "wydmre" that roughly 
    # represents the location.
    def _geohash(lat, lon, precision=6):
        # small geohash helper (optional metadata)
        if lat == 0 and lon == 0:
            return ""
        lat_interval = [-90.0, 90.0]
        lon_interval = [-180.0, 180.0]
        bits = [16, 8, 4, 2, 1]
        is_lon = True
        bit = 0
        ch = 0
        geohash_str = []

        for _ in range(precision * 5):
            if is_lon:
                mid = (lon_interval[0] + lon_interval[1]) / 2
                if lon > mid:
                    ch |= bits[bit]
                    lon_interval[0] = mid
                else:
                    lon_interval[1] = mid
            else:
                mid = (lat_interval[0] + lat_interval[1]) / 2
                if lat > mid:
                    ch |= bits[bit]
                    lat_interval[0] = mid
                else:
                    lat_interval[1] = mid
            is_lon = not is_lon
            bit += 1
            if bit == 5:
                geohash_str.append(_BASE32[ch])
                bit = 0
                ch = 0

        return "".join(geohash_str)

    # ----- Parse GeoJSON features into Documents -----
    documents = []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            gj = json.load(f)
        features = gj.get('features', []) if isinstance(gj, dict) else []

        # Loop over each feature and build Documents (props dictionary)
        for i, feat in enumerate(features):
            props = feat.get('properties', {}) if isinstance(feat, dict) else {}

            # 1) Extract si / gu / dong names and codes
            si_name   = _norm_str(
                props.get("city_name")
                or props.get("si")
                or props.get("location_si")
            )
            gu_name   = _norm_str(
                props.get("district_name")
                or props.get("gu")
                or props.get("location_gu")
            )
            dong_name = _norm_str(
                props.get("area_name")
                or props.get("dong")
                or props.get("location_dong")
            )

            si_code   = _norm_str(props.get("city_code") or props.get("si_code"))
            gu_code   = _norm_str(props.get("district_code") or props.get("gu_code"))
            area_code = _norm_str(
                props.get("area_code")
                or props.get("dong_code")
                or props.get("code")
            )
            
            # 2) Extract numeric fields (population, density, etc.)
            population     = _safe_float(props.get("total_population"), 0)
            density        = _safe_float(props.get("population_density"), 0)
            elderly_index  = _safe_float(
                props.get("elderly_index") or props.get("aging_index"), 0
            )
            avg_age        = _safe_float(props.get("average_age"), 0)
            area_sqkm      = _safe_float(props.get("area_sqkm"), 0)
            centroid_lat   = _safe_float(props.get("centroid_lat"), 0)
            centroid_lon   = _safe_float(props.get("centroid_lon"), 0)

            # 3) Calculate risk labels from numbers
            # Density Risk (인구밀도 위험도)       
            if density > 15000:
                density_risk = "매우 높음"
            elif density > 10000:
                density_risk = "높음"
            elif density > 5000:
                density_risk = "중간"
            else:
                density_risk = "낮음"
                
            # Population scale (인구 규모)
            if population > 500000:
                pop_risk = "대규모"
            elif population > 100000:
                pop_risk = "중규모"
            elif population > 10000:
                pop_risk = "소규모"
            else:
                pop_risk = "미소규모"

            # Elderly concern (고령 인구 관리 필요 여부)
            elderly_concern = "예" if elderly_index > 100 else "아니오"

            # Evacuation difficulty -> 대피 난이도 (예시: 밀도 기반)
            if density > 15000:
                evacuation_difficulty = "장시간 소요 (고밀도)"
            elif density > 10000:
                evacuation_difficulty = "중간 (중밀도)"
            else:
                evacuation_difficulty = "짧음 (저밀도)"

            # 4) Combine original props + GIS_인구 info into ONE JSON object
            combined_props = dict(props)  # shallow copy so we don't mutate original

            # Add extra GIS_인구 fields into the same JSON
            combined_props["인구밀도_위험도"] = density_risk
            combined_props["인구_규모"] = pop_risk
            combined_props["고령_인구_관리_필요"] = elderly_concern
            combined_props["예상_대피_난이도"] = evacuation_difficulty

            # Single JSON string as page_content
            full_page_content = json.dumps(combined_props, ensure_ascii=False, indent=2)

            # 5) metadata: keep old keys, just add extras
            geoh = _geohash(centroid_lat, centroid_lon, precision=6)

            metadata = {
                "source": file_path,
                "feature_index": i,
                "si_name": si_name,
                "gu_name": gu_name,
                "dong_name": dong_name,
                "si_code": si_code,
                "gu_code": gu_code,
                "area_code": area_code,
                "population": int(population),
                "population_density": density,
                "average_age": avg_age,
                "elderly_index": elderly_index,
                "area_sqkm": area_sqkm,
                "centroid_lat": centroid_lat,
                "centroid_lon": centroid_lon,
                "geohash": geoh,
                "density_risk": density_risk,
                "pop_risk": pop_risk,
                "elderly_concern": elderly_concern,
                "evacuation_difficulty": evacuation_difficulty,
            }

            documents.append(
                Document(
                    page_content=full_page_content,
                    metadata=metadata
                )
            )

    except Exception as e:
        print(f"Error loading GeoJSON file {file_path}: {e}")

    return documents


def load_all_documents_to_list(directory_path):
    documents = []

    all_documents = glob.glob(os.path.join(directory_path,"*"))

    for file_path in all_documents:
        try:
            if os.path.isdir(file_path):#폴더는 건너뛰기
                print(f"폴더 건너뜀: {file_path}")
                continue

            elif file_path.endswith(".pdf"): #확장자 pdf로 끝나는 경우 
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        #일반 텍스트 추출 
                        text = page.extract_text() or ""
                        
                        #표 데이터 추출
                        tables = page.extract_tables()
                        table_text = ""
                        if tables:
                            for table in tables:
                                # 추출된 표 데이터를 문자열로 변환-
                                table_text += "\n\n--- TABLE START ---\n"
                                for row in table:
                                    table_text += " | ".join(map(str, row)) + "\n"
                                table_text += "--- TABLE END ---\n\n"
                        
                        # 텍스트와 표 텍스트를 합쳐 Document 생성
                        full_content = text + table_text
                        documents.append(Document(page_content=full_content, metadata={"source": file_path, "page": page_num}))
                        #print(f" 성공: PDF 파일 처리 완료 -> {file_path}, 총 {len(reader.pages)} 페이지 추출 ",flush=True)

            elif file_path.endswith(".docx"): #docx로 끝나는 경우
                doc = docx.Document(file_path)
                full_text = "\n".join([para.text for para in doc.paragraphs])
                documents.append(Document(page_content=full_text,metadata={"source": file_path}))

            elif file_path.endswith(".hwpx"):
                full_text=""
                with zipfile.ZipFile(file_path,'r') as z:
                    section_xml_files = [f for f in z.namelist() if f.startswith('Contents/section') and f.endswith('.xml')]
                    for section_file in sorted(section_xml_files):
                        xml_content = z.read(section_file)
                        root = ET.fromstring(xml_content)
                        for text_element in root.iter('{http://www.hancom.co.kr/hwpml/2011/paragraph}t'):
                            if text_element.text:
                                full_text += text_element.text + "\n"
                documents.append(Document(page_content=full_text, metadata={"source":file_path}))

            elif file_path.endswith(".txt"): # 확장자가 .txt로 끝나는 경우
                # 'r' 모드(읽기 전용), 'utf-8' 인코딩으로 파일을 엽니다.
                with open(file_path, 'r', encoding='utf-8') as f:
                    raw_text = f.read() 
                    pattern = r'^.*사진.*$\n?' 
                    full_text = re.sub(pattern, '', raw_text, flags=re.MULTILINE)
                documents.append(
                    Document(
                        page_content=full_text, 
                        metadata={"source": os.path.basename(file_path)}
                    )
                )
                #print(f" 성공: TXT 파일 처리 완료 -> {file_path}, 텍스트 길이: {len(full_text)}", flush=True)
            
            elif file_path.endswith(".xlsx"): #excel(모바일상황실 접수 현황 검토중_240705_16시40분까지 (1).xlxs 파일만 적용)
                """ 읽어올 데이터 
                1. 1행 G열
                2. 2행~끝까지 F(유형),G(세부사항),H(주소),O(최종소요시간),V(첫 조치 소요시간) 행 
                """
                #1행 G열 데이터 읽기 (딕셔너리 형태로 받아옴)
                df1 = pd.read_excel(file_path, sheet_name=0,header=None, nrows=1)  # 첫 번째 시트만 읽기, 첫행 데이터로 취급, 1행의 G열만 필요 : 1행만 읽기 
                g_value = df1.iloc[0,6]
                parse_data_g_value = parse_multiline_cell(g_value)
                header_content = ", ".join([f"{k}: {v}" for k, v in parse_data_g_value.items()])
                doc = Document(
                    page_content=f"문서 요약 정보: {header_content}",
                    metadata={"source": os.path.basename(file_path), "section": "header"}
                )
                documents.append(doc)

                #2행부터 끝까지 읽어옴
                df2= pd.read_excel(file_path, sheet_name=0, header=1, usecols="F,G,H,O,V")
                for index, row in df2.iterrows():
                    row_content = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    doc = Document(
                        page_content=row_content,
                        metadata={
                            "source": os.path.basename(file_path),
                            "row": index + 3 # 엑셀 실제 행 번호와 맞춤
                        }
                    )
                    documents.append(doc)

        except Exception as e:
            # 해당 파일은 건너뛰고 오류 메시지를 출력
            print(f"Error processing {file_path}: {e}")
    #print(documents)
    return documents

def parse_multiline_cell(cell_text):
    """
    "키: 값" 형태의 여러 줄 문자열을 딕셔너리로 변환
    """
    if not isinstance(cell_text, str):
        return {}
        
    data_dict = {}
    lines = cell_text.splitlines() #\n 기준 나누기
    
    for line in lines:
        # 2. 각 줄을 ':' 기준으로 -> ['유형', '연쇄사항']
        if ':' in line:
            key, value = line.split(':', 1)
            data_dict[key.strip()] = value.strip()
            
    return data_dict


def build_vectorstores():
    """문서 분할 + 벡터DB 생성"""
    law_docs = load_all_documents_to_list("Dataset/관련법령")
    law_flooding_docs = load_all_documents_to_list("Dataset/관련법령/법령_풍수해")
    law_blackout_docs = load_all_documents_to_list("Dataset/관련법령/법령_정전")
    manual_docs = load_all_documents_to_list("Dataset/매뉴얼")
    basic_docs = load_all_documents_to_list("Dataset/기본데이터")
    past_docs = load_all_documents_to_list("Dataset/과거재난데이터")
    pop_docs = _load_geojson_as_docs("Location_Population_Data/location_query_result.geojson")

    #law_docs, manual_docs = load_all_documents_to_list("Dataset_for_test/과거재난데이터"), load_all_documents_to_list("Dataset_for_test/매뉴얼")
    splitter = CharacterTextSplitter(chunk_size=200, chunk_overlap=0)

    law_splits = splitter.split_documents(law_docs)
    law_flooding_splits = splitter.split_documents(law_flooding_docs)
    law_blackout_splits = splitter.split_documents(law_blackout_docs)
    manual_splits = splitter.split_documents(manual_docs)
    basic_splits = splitter.split_documents(basic_docs)
    pop_splits = splitter.split_documents(pop_docs)
    past_splits = splitter.split_documents(past_docs)

    embeddings = load_embeddings()

    vectordb_law = FAISS.from_documents(law_splits, embeddings)
    vectordb_flooding_law = FAISS.from_documents(law_flooding_splits, embeddings)
    vectordb_blackout_law = FAISS.from_documents(law_blackout_splits, embeddings)
    vectordb_manual = FAISS.from_documents(manual_splits, embeddings)
    vectordb_basic = FAISS.from_documents(basic_splits, embeddings)
    vectordb_population = FAISS.from_documents(pop_splits, embeddings)
    vectordb_past = FAISS.from_documents(past_splits, embeddings)

    return (
        vectordb_law,
        vectordb_flooding_law,
        vectordb_blackout_law,
        vectordb_manual,
        vectordb_basic,
        vectordb_population,
        vectordb_past,
    )
